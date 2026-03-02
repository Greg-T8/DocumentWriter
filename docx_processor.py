# -------------------------------------------------------------------------
# Program: docx_processor.py
# Description: Extract sections/images from a Word document and insert
#              AI-generated commentary back into the document.
# Context: DocumentWriter project - GitHub Models integration
# Author: Greg Tate
# -------------------------------------------------------------------------

"""
Processes Word (.docx) documents: extracts section outlines and embedded
images for AI analysis, and inserts generated commentary paragraphs back
into the document.

Usage:
    python docx_processor.py extract <docx_path> <output_json>
    python docx_processor.py insert  <docx_path> <commentary_json> <output_docx>
"""

#region IMPORTS
import sys
import json
import base64
import re
from pathlib import Path
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor
from PIL import Image
#endregion


#region MAIN WORKFLOW
def main() -> None:
    """
    Main entry point - dispatches to extract or insert based on CLI args.
    """
    validate_argument(sys.argv)

    # Determine which command to run
    command = sys.argv[1].lower()

    if command == "extract":
        run_extract(sys.argv[2], sys.argv[3])
    elif command == "insert":
        run_insert(sys.argv[2], sys.argv[3], sys.argv[4])
#endregion


#region EXTRACT FUNCTIONS
def run_extract(
    docx_path: str,
    output_json: str
) -> None:
    """
    Extract sections and images from a Word document and write JSON output.

    Args:
        docx_path: Path to the input .docx file
        output_json: Path to write the extracted JSON data
    """
    doc = Document(docx_path)

    # Extract all embedded images from the document
    image_map = extract_image(doc)

    # Walk paragraphs and build section list
    sections = build_section(doc, image_map)

    # Write the extracted data to JSON
    output = {
        "source_document": str(Path(docx_path).resolve()),
        "sections": sections
    }

    with open(output_json, "w", encoding="utf-8") as f:
        json.dump(output, f, indent=2, ensure_ascii=False)

    print(f"Extracted {len(sections)} sections to {output_json}")


def extract_image(doc: Document) -> dict:
    """
    Extract all images from a document's media relationships.

    Returns:
        dict: Mapping of relationship ID to base64-encoded image data
    """
    image_map = {}

    for rel_id, rel in doc.part.rels.items():
        # Only process image relationships
        if "image" in rel.reltype:
            image_data = rel.target_part.blob
            image_b64 = base64.b64encode(image_data).decode("utf-8")

            # Detect MIME type from the image bytes
            mime = detect_mime_type(image_data)

            image_map[rel_id] = {
                "base64": image_b64,
                "mime_type": mime,
                "size_bytes": len(image_data)
            }

    return image_map


def detect_mime_type(image_data: bytes) -> str:
    """
    Detect the MIME type of an image from its binary data.

    Args:
        image_data: Raw image bytes

    Returns:
        str: MIME type string (e.g., 'image/png')
    """
    try:
        img = Image.open(BytesIO(image_data))
        fmt = (img.format or "png").lower()
        mime_map = {
            "png": "image/png",
            "jpeg": "image/jpeg",
            "gif": "image/gif",
            "bmp": "image/bmp",
            "tiff": "image/tiff",
            "webp": "image/webp",
        }
        return mime_map.get(fmt, f"image/{fmt}")
    except Exception:
        return "image/png"


def build_section(
    doc: Document,
    image_map: dict
) -> list:
    """
    Walk document paragraphs and group them into sections by heading.

    Args:
        doc: The loaded Document object
        image_map: Mapping of relationship IDs to image data

    Returns:
        list: List of section dictionaries with headings, text, and images
    """
    sections = []
    current_section = None

    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""

        # Detect heading paragraphs to start a new section
        if style_name.startswith("Heading"):
            if current_section is not None:
                sections.append(current_section)

            # Parse heading level from style name
            level = extract_heading_level(style_name)

            current_section = {
                "heading": paragraph.text.strip(),
                "heading_level": level,
                "text_content": [],
                "images": [],
                "paragraph_index_start": get_paragraph_index(
                    doc, paragraph
                )
            }
        elif current_section is not None:
            # Accumulate body text for the current section
            if paragraph.text.strip():
                current_section["text_content"].append(
                    paragraph.text.strip()
                )

            # Check for inline images in this paragraph
            images = extract_paragraph_image(paragraph, image_map)
            if images:
                current_section["images"].extend(images)

    # Append the last section if present
    if current_section is not None:
        sections.append(current_section)

    # Handle documents with no headings - treat entire body as one section
    if not sections:
        sections.append(
            create_default_section(doc, image_map)
        )

    return sections


def extract_heading_level(style_name: str) -> int:
    """
    Parse the heading level number from a Word style name.

    Args:
        style_name: The paragraph style name (e.g., 'Heading 2')

    Returns:
        int: The heading level, or 0 if not parsable
    """
    match = re.search(r"(\d+)", style_name)
    if match:
        return int(match.group(1))
    return 0


def get_paragraph_index(
    doc: Document,
    target_paragraph
) -> int:
    """
    Find the index of a paragraph within the document body.

    Args:
        doc: The Document object
        target_paragraph: The paragraph to locate

    Returns:
        int: Zero-based index of the paragraph
    """
    for i, para in enumerate(doc.paragraphs):
        if para._element is target_paragraph._element:
            return i
    return -1


def extract_paragraph_image(
    paragraph,
    image_map: dict
) -> list:
    """
    Extract image references from inline runs within a paragraph.

    Args:
        paragraph: A python-docx Paragraph object
        image_map: Mapping of relationship IDs to image data

    Returns:
        list: List of image data dictionaries found in the paragraph
    """
    images = []

    # Namespace for drawing elements
    nsmap = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": (
            "http://schemas.openxmlformats.org/"
            "officeDocument/2006/relationships"
        ),
        "wp": (
            "http://schemas.openxmlformats.org/"
            "drawingml/2006/wordprocessingDrawing"
        ),
    }

    # Search for blip elements that reference embedded images
    for blip in paragraph._element.findall(
        ".//a:blip", nsmap
    ):
        embed_attr = blip.get(
            "{http://schemas.openxmlformats.org/"
            "officeDocument/2006/relationships}embed"
        )

        if embed_attr and embed_attr in image_map:
            images.append(image_map[embed_attr])

    return images


def create_default_section(
    doc: Document,
    image_map: dict
) -> dict:
    """
    Create a single section for documents that have no headings.

    Args:
        doc: The Document object
        image_map: Mapping of relationship IDs to image data

    Returns:
        dict: A section dictionary covering the entire document
    """
    all_text = [
        p.text.strip()
        for p in doc.paragraphs
        if p.text.strip()
    ]

    all_images = []
    for paragraph in doc.paragraphs:
        imgs = extract_paragraph_image(paragraph, image_map)
        if imgs:
            all_images.extend(imgs)

    return {
        "heading": "Document Content",
        "heading_level": 0,
        "text_content": all_text,
        "images": all_images,
        "paragraph_index_start": 0,
    }
#endregion


#region INSERT FUNCTIONS
def run_insert(
    docx_path: str,
    commentary_json: str,
    output_path: str
) -> None:
    """
    Insert AI-generated commentary into a copy of the Word document.

    Args:
        docx_path: Path to the original .docx file
        commentary_json: Path to JSON file containing commentary entries
        output_path: Path for the annotated output .docx file
    """
    doc = Document(docx_path)

    # Load the commentary data
    with open(commentary_json, "r", encoding="utf-8") as f:
        commentary_data = json.load(f)

    # Insert commentary after each matching section heading
    insert_count = insert_commentary(doc, commentary_data)

    # Save the annotated document
    doc.save(output_path)
    print(
        f"Inserted {insert_count} commentary blocks "
        f"into {output_path}"
    )


def insert_commentary(
    doc: Document,
    commentary_data: list
) -> int:
    """
    Insert formatted commentary paragraphs after section headings.

    Args:
        doc: The Document object to modify
        commentary_data: List of dicts with 'heading' and 'commentary' keys

    Returns:
        int: Number of commentary blocks inserted
    """
    insert_count = 0

    # Track offset as we insert new paragraphs
    offset = 0

    for entry in commentary_data:
        heading_text = entry.get("heading", "")
        commentary_text = entry.get("commentary", "")

        # Skip entries with no commentary
        if not commentary_text.strip():
            continue

        # Find the heading paragraph in the document
        target_index = find_heading_paragraph(doc, heading_text)

        if target_index < 0:
            continue

        # Calculate the adjusted index accounting for prior insertions
        adjusted_index = target_index + offset

        # Find the end of the section body to insert after it
        insert_after = find_section_end(doc, adjusted_index)

        # Insert the commentary paragraph
        add_commentary_paragraph(doc, insert_after, commentary_text)
        offset += 1
        insert_count += 1

    return insert_count


def find_heading_paragraph(
    doc: Document,
    heading_text: str
) -> int:
    """
    Find the index of a paragraph matching the given heading text.

    Args:
        doc: The Document object
        heading_text: Text content of the heading to find

    Returns:
        int: Index of the matching paragraph, or -1 if not found
    """
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""

        if (
            style_name.startswith("Heading")
            and para.text.strip() == heading_text.strip()
        ):
            return i

    return -1


def find_section_end(
    doc: Document,
    heading_index: int
) -> int:
    """
    Find the last paragraph index belonging to a section.

    Scans forward from the heading until the next heading or document end.

    Args:
        doc: The Document object
        heading_index: Index of the section heading paragraph

    Returns:
        int: Index of the last paragraph in the section
    """
    last_index = heading_index

    for i in range(heading_index + 1, len(doc.paragraphs)):
        style_name = doc.paragraphs[i].style.name if doc.paragraphs[i].style else ""

        # Stop at the next heading
        if style_name.startswith("Heading"):
            break

        last_index = i

    return last_index


def add_commentary_paragraph(
    doc: Document,
    after_index: int,
    commentary_text: str
) -> None:
    """
    Insert a visually distinct commentary paragraph after the given index.

    Args:
        doc: The Document object
        after_index: Paragraph index to insert after
        commentary_text: The commentary text to insert
    """
    # Build the new paragraph element
    target_element = doc.paragraphs[after_index]._element
    new_para = doc.add_paragraph()

    # Style the commentary paragraph
    new_para.style = doc.styles["Normal"]
    new_para.paragraph_format.space_before = Pt(6)
    new_para.paragraph_format.space_after = Pt(6)

    # Add a bold label prefix
    label_run = new_para.add_run("[AI Commentary] ")
    label_run.bold = True
    label_run.font.color.rgb = RGBColor(0x00, 0x51, 0x8A)
    label_run.font.size = Pt(10)

    # Add the commentary text
    text_run = new_para.add_run(commentary_text)
    text_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    text_run.font.size = Pt(10)
    text_run.italic = True

    # Move the new paragraph element to the correct position
    new_element = new_para._element
    new_element.getparent().remove(new_element)
    target_element.addnext(new_element)
#endregion


#region VALIDATION
def validate_argument(argv: list) -> None:
    """
    Validate command-line arguments and print usage if incorrect.

    Args:
        argv: The sys.argv list
    """
    if len(argv) < 2:
        show_usage()
        sys.exit(1)

    command = argv[1].lower()

    # Validate extract command requires 2 additional args
    if command == "extract" and len(argv) != 4:
        print("Error: 'extract' requires <docx_path> <output_json>")
        show_usage()
        sys.exit(1)

    # Validate insert command requires 3 additional args
    if command == "insert" and len(argv) != 5:
        print(
            "Error: 'insert' requires "
            "<docx_path> <commentary_json> <output_docx>"
        )
        show_usage()
        sys.exit(1)

    # Validate command name
    if command not in ("extract", "insert"):
        print(f"Error: Unknown command '{command}'")
        show_usage()
        sys.exit(1)


def show_usage() -> None:
    """
    Print usage instructions to stderr.
    """
    print(
        "Usage:\n"
        "  python docx_processor.py extract "
        "<docx_path> <output_json>\n"
        "  python docx_processor.py insert  "
        "<docx_path> <commentary_json> <output_docx>",
        file=sys.stderr,
    )
#endregion


# -------------------------------------------------------------------------
# Script Entry Point
# -------------------------------------------------------------------------

if __name__ == "__main__":
    main()
