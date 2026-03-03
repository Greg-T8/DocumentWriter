# -------------------------------------------------------------------------
# Program: docx_processor.py
# Description: Extract text sections from a Word document and revise
#              the document based on AI analysis.
# Context: DocumentWriter project - GitHub Models integration
# Author: Greg Tate
# -------------------------------------------------------------------------

"""
Processes Word (.docx) documents: extracts text section outlines for AI
analysis, and revises the document based on AI-generated suggestions.

Usage:
    python docx_processor.py extract <docx_path> <output_json>
    python docx_processor.py revise  <docx_path> <revision_json> <output_docx>
"""

#region IMPORTS
import sys
import json
import re
from pathlib import Path

from docx import Document
#endregion


#region MAIN WORKFLOW
def main() -> None:
    """
    Main entry point - dispatches to extract or revise based on CLI args.
    """
    # Determine which command to run
    validate_argument(sys.argv)

    command = sys.argv[1].lower()

    if command == "extract":
        run_extract(sys.argv[2], sys.argv[3])
    elif command in ("insert", "revise"):
        run_revise(sys.argv[2], sys.argv[3], sys.argv[4])
#endregion


#region EXTRACT FUNCTIONS
def run_extract(
    docx_path: str,
    output_json: str
) -> None:
    """
    Extract text sections from a Word document and write JSON output.

    Args:
        docx_path: Path to the input .docx file
        output_json: Path to write the extracted JSON data
    """
    doc = Document(docx_path)

    # Walk paragraphs and build section list
    sections = build_section(doc)

    # Write the extracted data to JSON
    output = {
        "source_document": str(Path(docx_path).resolve()),
        "sections": sections
    }

    with open(output_json, "w", encoding="utf-8") as f:
        json.dump(output, f, indent=2, ensure_ascii=False)

    print(f"Extracted {len(sections)} sections to {output_json}")

def build_section(
    doc: Document
) -> list:
    """
    Walk document paragraphs and group them into sections by heading.

    Args:
        doc: The loaded Document object

    Returns:
        list: List of section dictionaries with headings and text
    """
    sections = []
    current_section = None

    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""

        # Detect heading paragraphs to start a new section
        if style_name.startswith("Heading"):
            if current_section is not None:
                sections.append(current_section)

            # Parse heading level from style name
            level = extract_heading_level(style_name)

            current_section = {
                "heading": paragraph.text.strip(),
                "heading_level": level,
                "text_content": [],
                "paragraph_index_start": get_paragraph_index(
                    doc, paragraph
                )
            }
        elif current_section is not None:
            # Accumulate body text for the current section
            if paragraph.text.strip():
                current_section["text_content"].append(
                    paragraph.text.strip()
                )

    # Append the last section if present
    if current_section is not None:
        sections.append(current_section)

    # Handle documents with no headings - treat entire body as one section
    if not sections:
        sections.append(
            create_default_section(doc)
        )

    return sections


def extract_heading_level(style_name: str) -> int:
    """
    Parse the heading level number from a Word style name.

    Args:
        style_name: The paragraph style name (e.g., 'Heading 2')

    Returns:
        int: The heading level, or 0 if not parsable
    """
    match = re.search(r"(\d+)", style_name)
    if match:
        return int(match.group(1))
    return 0


def get_paragraph_index(
    doc: Document,
    target_paragraph
) -> int:
    """
    Find the index of a paragraph within the document body.

    Args:
        doc: The Document object
        target_paragraph: The paragraph to locate

    Returns:
        int: Zero-based index of the paragraph
    """
    for i, para in enumerate(doc.paragraphs):
        if para._element is target_paragraph._element:
            return i
    return -1

def create_default_section(
    doc: Document
) -> dict:
    """
    Create a single section for documents that have no headings.

    Args:
        doc: The Document object

    Returns:
        dict: A section dictionary covering the entire document
    """
    all_text = [
        p.text.strip()
        for p in doc.paragraphs
        if p.text.strip()
    ]

    return {
        "heading": "Document Content",
        "heading_level": 0,
        "text_content": all_text,
        "paragraph_index_start": 0,
    }
#endregion


#region INSERT FUNCTIONS
def run_revise(
    docx_path: str,
    revision_json: str,
    output_path: str
) -> None:
    """
    Rewrite section body prose in a copy of the Word document.

    Args:
        docx_path: Path to the original .docx file
        revision_json: Path to JSON file containing rewritten section text
        output_path: Path for the revised output .docx file
    """
    source_path = Path(docx_path).resolve(strict=False)
    target_path = Path(output_path).resolve(strict=False)

    # Ensure the source document is never overwritten
    if source_path == target_path:
        raise ValueError(
            "Output document path must be different from input document path."
        )

    doc = Document(docx_path)

    # Load the section revision data
    with open(revision_json, "r", encoding="utf-8") as f:
        revision_data = json.load(f)

    # Rewrite matching section body paragraphs with revised prose
    revised_count = rewrite_section_body(doc, revision_data)

    # Save the revised document
    doc.save(output_path)
    print(
        f"Rewrote {revised_count} section bodies "
        f"into {output_path}"
    )


def rewrite_section_body(
    doc: Document,
    revision_data: list
) -> int:
    """
    Replace section body paragraphs with revised prose.

    Args:
        doc: The Document object to modify
        revision_data: List of dicts with 'heading' and 'revised_text' keys

    Returns:
        int: Number of section bodies rewritten
    """
    revised_count = 0

    for entry in revision_data:
        heading_text = entry.get("heading", "")
        revised_text = get_revised_text(entry)

        # Skip entries with no rewritten text
        if not revised_text.strip():
            continue

        # Find the heading paragraph in the document
        target_index = find_heading_paragraph(doc, heading_text)

        if target_index < 0:
            continue

        # Rewrite paragraph body for this section
        rewrite_section_by_index(doc, target_index, revised_text)
        revised_count += 1

    return revised_count


def get_revised_text(entry: dict) -> str:
    """
    Resolve revised text from a revision entry.

    Accepts either a preferred 'revised_text' key or a legacy
    'commentary' key for compatibility.

    Args:
        entry: The revision entry object

    Returns:
        str: Rewritten prose text
    """
    revised_text = entry.get("revised_text", "")
    if isinstance(revised_text, str) and revised_text.strip():
        return revised_text

    legacy_text = entry.get("commentary", "")
    if isinstance(legacy_text, str):
        return legacy_text

    return ""


def find_heading_paragraph(
    doc: Document,
    heading_text: str
) -> int:
    """
    Find the index of a paragraph matching the given heading text.

    Args:
        doc: The Document object
        heading_text: Text content of the heading to find

    Returns:
        int: Index of the matching paragraph, or -1 if not found
    """
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""

        if (
            style_name.startswith("Heading")
            and para.text.strip() == heading_text.strip()
        ):
            return i

    return -1


def get_section_body_indexes(
    doc: Document,
    heading_index: int
) -> list:
    """
    Collect paragraph indexes for the body belonging to a section.

    Args:
        doc: The Document object
        heading_index: Index of the section heading paragraph

    Returns:
        list: Paragraph indexes between this heading and the next heading
    """
    body_indexes = []

    for index in range(heading_index + 1, len(doc.paragraphs)):
        style_name = doc.paragraphs[index].style.name if doc.paragraphs[index].style else ""

        # Stop at the next heading
        if style_name.startswith("Heading"):
            break

        body_indexes.append(index)

    return body_indexes


def rewrite_section_by_index(
    doc: Document,
    heading_index: int,
    revised_text: str
) -> None:
    """
    Replace text body paragraphs in a section with rewritten content.

    Args:
        doc: The Document object
        heading_index: Index of the section heading paragraph
        revised_text: Rewritten prose to insert for the section
    """
    body_indexes = get_section_body_indexes(doc, heading_index)
    paragraph_texts = split_revised_paragraphs(revised_text)
    body_paragraphs = [doc.paragraphs[index] for index in body_indexes]

    # Keep media paragraphs (screenshots/images) and only replace text prose
    media_paragraphs = [
        paragraph
        for paragraph in body_paragraphs
        if paragraph_contains_media(paragraph)
    ]
    text_paragraphs = [
        paragraph
        for paragraph in body_paragraphs
        if paragraph not in media_paragraphs
    ]

    # Nothing to rewrite when there is no text body and no revised prose
    if not text_paragraphs and not paragraph_texts:
        return

    # Choose insertion anchor so media keeps its original relative position
    insertion_before_element = None
    insertion_after_element = None

    if text_paragraphs:
        insertion_before_element = text_paragraphs[0]._element
    else:
        insertion_after_element = doc.paragraphs[heading_index]._element

    # Insert rewritten paragraphs before removing old text paragraphs
    for paragraph_text in paragraph_texts:
        new_paragraph = doc.add_paragraph(paragraph_text)
        new_paragraph.style = doc.styles["Normal"]

        new_element = new_paragraph._element
        new_element.getparent().remove(new_element)

        if insertion_before_element is not None:
            insertion_before_element.addprevious(new_element)
        else:
            insertion_after_element.addnext(new_element)
            insertion_after_element = new_element

    # Remove only old text paragraphs from bottom to top
    for paragraph in reversed(text_paragraphs):
        remove_paragraph(paragraph)


def paragraph_contains_media(paragraph) -> bool:
    """
    Check whether a paragraph contains embedded media content.

    Args:
        paragraph: Paragraph object from python-docx

    Returns:
        bool: True when paragraph contains drawings, pictures, or objects
    """
    element = paragraph._element

    # Detect inline/floating images and legacy picture/object elements
    media_nodes = element.xpath(
        ".//*[local-name()='drawing' or local-name()='pict' or local-name()='object']"
    )

    return len(media_nodes) > 0


def remove_paragraph(paragraph) -> None:
    """
    Remove a paragraph from the document body.

    Args:
        paragraph: The paragraph object to remove
    """
    element = paragraph._element
    parent = element.getparent()

    if parent is not None:
        parent.remove(element)


def strip_markdown(text: str) -> str:
    """
    Remove common markdown formatting characters from model output.

    Handles bold/italic markers, inline code, links, heading prefixes,
    and bullet list markers so that plain text is inserted into Word.

    Args:
        text: Raw model output that may contain markdown syntax

    Returns:
        str: Plain text with markdown syntax removed
    """
    # Remove bold markers (**text** and __text__)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'__(.*?)__', r'\1', text)

    # Remove italic markers (*text* and _text_)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'_(.*?)_', r'\1', text)

    # Remove inline code backticks
    text = re.sub(r'`(.*?)`', r'\1', text)

    # Remove markdown links, keeping the display text
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)

    # Remove heading marker prefixes (e.g. ## Heading)
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)

    # Remove bullet list markers (-, *, +)
    text = re.sub(r'^[\s]*[-*+]\s+', '', text, flags=re.MULTILINE)

    # Remove numbered list markers (e.g. 1. , 2. )
    text = re.sub(r'^\s*\d+\.\s+', '', text, flags=re.MULTILINE)

    return text


def split_revised_paragraphs(
    revised_text: str
) -> list:
    """
    Split model output into document paragraph blocks.

    Strips markdown formatting before splitting so that plain text
    is inserted into the Word document.

    Args:
        revised_text: Full rewritten text from the model

    Returns:
        list: Non-empty paragraph strings
    """
    cleaned = revised_text.replace("\r\n", "\n").strip()
    if not cleaned:
        return []

    # Remove markdown syntax before inserting into Word
    cleaned = strip_markdown(cleaned)

    blocks = re.split(r"\n\s*\n", cleaned)
    return [block.strip() for block in blocks if block.strip()]
#endregion


#region VALIDATION
def validate_argument(argv: list) -> None:
    """
    Validate command-line arguments and print usage if incorrect.

    Args:
        argv: The sys.argv list
    """
    if len(argv) < 2:
        show_usage()
        sys.exit(1)

    command = argv[1].lower()

    # Validate extract command requires 2 additional args
    if command == "extract" and len(argv) != 4:
        print("Error: 'extract' requires <docx_path> <output_json>")
        show_usage()
        sys.exit(1)

    # Validate insert/revise command requires 3 additional args
    if command in ("insert", "revise") and len(argv) != 5:
        print(
            "Error: 'insert'/'revise' requires "
            "<docx_path> <revision_json> <output_docx>"
        )
        show_usage()
        sys.exit(1)

    # Validate command name
    if command not in ("extract", "insert", "revise"):
        print(f"Error: Unknown command '{command}'")
        show_usage()
        sys.exit(1)


def show_usage() -> None:
    """
    Print usage instructions to stderr.
    """
    print(
        "Usage:\n"
        "  python docx_processor.py extract "
        "<docx_path> <output_json>\n"
        "  python docx_processor.py revise  "
        "<docx_path> <revision_json> <output_docx>",
        file=sys.stderr,
    )
#endregion


# -------------------------------------------------------------------------
# Script Entry Point
# -------------------------------------------------------------------------

if __name__ == "__main__":
    main()
